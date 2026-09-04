"""Render a Collaborative Planning Guide (the JSON produced by app.ai) into an
editable Word document, mirroring the district's planning-guide format."""
import io
import os

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

# School logo shown at the top of every document, if the file is present.
# Set AVOCADO_LOGO to an absolute path, or drop it at app/data/school_logo.png.
_LOGO_PATHS = [
    os.environ.get("AVOCADO_LOGO", ""),
    os.path.join(os.path.dirname(__file__), "data", "school_logo.png"),
]


def _add_logo(doc):
    for p in _LOGO_PATHS:
        if p and os.path.exists(p):
            try:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run().add_picture(p, width=Inches(2.2))
                return
            except Exception:
                return


def _heading(doc, text, size, color=(0x38, 0x60, 0x1F), bold=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    return p


def _label(doc, label, value):
    if not value:
        return
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    p.add_run(str(value))


def _bullets(doc, items, style="List Bullet"):
    for it in items or []:
        doc.add_paragraph(str(it), style=style)


def _misconception_table(doc, rows, code_col=False):
    """Render the 3-column Misconception | Example Error | Correction Strategy
    table (with an optional leading Benchmark column for the topic-level view)."""
    if not rows:
        return
    headers = (["Benchmark"] if code_col else []) + \
        ["Misconception", "Example Error", "Correction Strategy"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for m in rows:
        cells = table.add_row().cells
        vals = ([m.get("code", "")] if code_col else []) + [
            m.get("misconception", ""), m.get("example", ""), m.get("fix", "")]
        for i, v in enumerate(vals):
            cells[i].text = str(v)


def _exit_ticket_text(et) -> str:
    if isinstance(et, dict):
        prob, ans = et.get("problem", ""), et.get("answer", "")
        return f"{prob}  →  {ans}" if ans else prob
    return str(et or "")


def _phase(doc, header, phase):
    """Render one fully-scripted ACES phase: the problem, the teacher's spoken
    script, what the teacher does, the Concrete/Pictorial/Abstract for that
    problem, CUBS (Solo), and the look-fors. Tolerates the legacy string shape."""
    if not phase:
        return
    _heading(doc, header, 11)
    if isinstance(phase, str):
        doc.add_paragraph(phase)
        return
    if phase.get("strategy"):
        _label(doc, "Strategy we are modeling", phase.get("strategy"))
    if phase.get("connect"):
        _label(doc, "Connect to what we just modeled", phase.get("connect"))
    if phase.get("structure"):
        _label(doc, "Collaborative strategy", phase.get("structure"))
        _label(doc, "Each partner/group role", phase.get("roles"))
    _label(doc, "Problem worked", phase.get("problem"))
    say = phase.get("say")
    if say:
        p = doc.add_paragraph()
        p.add_run("Teacher says:").bold = True
        for line in (say if isinstance(say, list) else [say]):
            doc.add_paragraph(f"“{line}”", style="List Bullet")
    questions = phase.get("questions")
    if questions:
        p = doc.add_paragraph()
        p.add_run("Ask these questions:").bold = True
        for q in (questions if isinstance(questions, list) else [questions]):
            doc.add_paragraph(str(q), style="List Bullet")
    _label(doc, "Teacher does", phase.get("do"))
    if phase.get("check"):
        _label(doc, "Check for understanding", phase.get("check"))
    _label(doc, "Concrete (manipulative)", phase.get("concrete"))
    _label(doc, "Pictorial (drawing)", phase.get("pictorial"))
    _label(doc, "Abstract (equation)", phase.get("abstract"))
    _label(doc, "CUBS on this problem", phase.get("cubs"))
    _label(doc, "Look for", phase.get("look_for"))


def guide_to_docx(guide: dict) -> bytes:
    doc = Document()
    _add_logo(doc)
    _heading(doc, guide.get("title", "Collaborative Planning Guide"), 16)
    meta = doc.add_paragraph()
    m = meta.add_run(
        f"Grade {guide.get('grade_level','')} {guide.get('subject','')}  ·  "
        + ("AI-generated draft" if guide.get("ai_generated") else "Draft")
    )
    m.italic = True
    m.font.size = Pt(9)

    # Quick Facts
    qf = guide.get("quick_facts", {})
    if qf:
        _heading(doc, "Quick Facts", 13)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        rows = [
            ("Time Frame", qf.get("time_frame", "")),
            ("Assessment Date", qf.get("assessment_date", "")),
            ("Topic Focus", qf.get("topic_focus", "")),
            ("Key Benchmarks", ", ".join(qf.get("key_benchmarks", []))),
            ("ALD Focus", qf.get("ald_focus", "")),
            ("MTR Practices", " · ".join(qf.get("mtr_practices", []))),
            ("Materials", ", ".join(qf.get("materials", []))),
        ]
        for k, v in rows:
            if not v:
                continue
            cells = table.add_row().cells
            cells[0].text = k
            cells[1].text = str(v)

    _label(doc, "Learning Goal", guide.get("learning_goal"))
    if guide.get("success_criteria"):
        _heading(doc, "Success Criteria", 12)
        _bullets(doc, guide["success_criteria"])

    # Tier 2 academic vocabulary (this year's focus) — mined from the standards.
    if guide.get("tier2_vocabulary"):
        _heading(doc, "Power Words (academic vocabulary — the words in the "
                      "question stems)", 12)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(["Word", "What it means (kid words)"]):
            table.rows[0].cells[i].text = h
        for w in guide["tier2_vocabulary"]:
            cells = table.add_row().cells
            cells[0].text = str(w.get("word", ""))
            cells[1].text = str(w.get("meaning", ""))

    if guide.get("benchmark_clarifications"):
        _heading(doc, "Benchmark Clarifications", 12)
        for c in guide["benchmark_clarifications"]:
            _label(doc, c.get("code", ""), c.get("description", ""))
            _bullets(doc, c.get("clarifications", []))

    if guide.get("common_misconceptions"):
        _heading(doc, "Common Misconceptions", 12)
        _misconception_table(doc, guide["common_misconceptions"], code_col=True)

    # Lessons
    for L in guide.get("lessons", []):
        doc.add_page_break()
        _heading(doc, f"Lesson {L.get('code','')} — {L.get('title','')}", 14)
        _label(doc, "Benchmarks", ", ".join(L.get("benchmarks", [])))
        _label(doc, "Focus", L.get("focus"))
        _label(doc, "Learning Goal (Student-Friendly)", L.get("learning_goal"))
        if L.get("success_criteria"):
            _heading(doc, "Success Criteria", 11)
            _bullets(doc, L["success_criteria"])
        _label(doc, "Example", L.get("success_example"))
        _label(doc, "Benchmark Clarification", L.get("benchmark_clarification"))
        _label(doc, "Example", L.get("benchmark_example"))
        _label(doc, "Sentence Frame", L.get("sentence_frame"))
        if L.get("vocabulary") or L.get("vocabulary_integration"):
            _heading(doc, "Vocabulary (from the pacing guide)", 11)
            _label(doc, "Terms", ", ".join(L.get("vocabulary", [])))
            _label(doc, "How to integrate", L.get("vocabulary_integration"))
        if L.get("misconceptions"):
            _heading(doc, "Common Misconceptions & Fixes", 11)
            _misconception_table(doc, L["misconceptions"])
        # ACES gradual release: Assemble (I Do) -> Connect (We Do) -> Explore
        # (Y'all Do, collaborative pairs/groups of 4) -> Solo (You Do + CUBS).
        if any(L.get(k) for k in ("activate_prior_knowledge", "i_do", "we_do",
                                  "explore_yall_do", "you_do")):
            _heading(doc, "Teaching Strategy — ACES Gradual Release (Scripted)", 12)
            _label(doc, "Activate Prior Knowledge", L.get("activate_prior_knowledge"))
            _phase(doc, "ASSEMBLE · I Do (Teacher Models)", L.get("i_do"))
            _phase(doc, "CONNECT · We Do (Guided Practice)", L.get("we_do"))
            _phase(doc, "EXPLORE · Y'all Do (Collaborative — pairs or groups of 4)",
                   L.get("explore_yall_do"))
            _phase(doc, "SOLO · You Do (Independent Practice + CUBS)", L.get("you_do"))
            # Legacy top-level CUBS string only when You Do isn't a scripted phase.
            if not isinstance(L.get("you_do"), dict):
                _label(doc, "SOLO · Apply CUBS to the problem", L.get("cubs"))
        elif L.get("teaching_strategy"):
            _heading(doc, "Teaching Strategy (Step-by-Step)", 11)
            _bullets(doc, L["teaching_strategy"], style="List Number")
        # Lesson-level CPA only when phases aren't already scripted with their own
        # Concrete/Pictorial/Abstract (avoids duplicating it four times over).
        cpa = L.get("cpa", {})
        if not isinstance(L.get("i_do"), dict) and any(
                cpa.get(k) for k in ("concrete", "pictorial", "abstract")):
            _heading(doc, "CPA Model", 11)
            _label(doc, "Concrete", cpa.get("concrete"))
            _label(doc, "Pictorial", cpa.get("pictorial"))
            _label(doc, "Abstract", cpa.get("abstract"))
        # Achievement Level Descriptors — what each level looks like for this
        # benchmark, with Level 3 (on-grade, the goal) emphasized.
        ald = L.get("ald") or {}
        if any(ald.get(k) for k in ("level2", "level3", "level4", "level5")):
            _heading(doc, "Achievement Level Descriptors (ALD)", 11)
            rows = [("Level 2 — below", ald.get("level2")),
                    ("Level 3 — ON GRADE (goal)", ald.get("level3")),
                    ("Level 4 — above", ald.get("level4")),
                    ("Level 5 — mastery", ald.get("level5"))]
            for lbl, val in rows:
                if not val:
                    continue
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(f"{lbl}: ")
                r.bold = True
                if "ON GRADE" in lbl:
                    r.font.color.rgb = RGBColor(0x1F, 0x7A, 0x1F)
                p.add_run(str(val))
        l3 = L.get("level3_look_like") or {}
        if l3.get("problem"):
            _heading(doc, "What a Level 3 (On-Grade) Looks Like — This Lesson", 11)
            _label(doc, "Problem", l3.get("problem"))
            _label(doc, "Worked solution", l3.get("solution"))
            _label(doc, "Student explanation", l3.get("student_explanation"))
        else:
            _label(doc, "Level 3 Proficiency Example (student voice)",
                   L.get("level3_example"))
        br = L.get("book_reference")
        if isinstance(br, dict) and (br.get("lesson") or br.get("pages")):
            _heading(doc, "📖 In the book — where each part comes from", 11)
            _label(doc, "Book lesson", br.get("lesson"))
            _label(doc, "Pages", br.get("pages"))

            def _bp(label, key, pkey):
                val = br.get(key)
                if not val:
                    return
                pg = br.get(pkey)
                _label(doc, label, f"{val}" + (f"  (p. {pg})" if pg else ""))

            _bp("I Do — model (Modeling Real Life)", "model_example", "model_pages")
            _bp("We Do — guided (Try It / Show and Grow)", "guided_practice", "guided_pages")
            _bp("Solo/You Do (In-Class Practice)", "independent_practice", "independent_pages")
            _bp("Exit slip (In-Class Practice / Closure)", "exit_problem", "exit_pages")
            _bp("Level 3 target problem", "level3_problem", "level3_pages")
            _bp("Dig Deeper (stretch / enrichment)", "dig_deeper", "dig_deeper_pages")
            _label(doc, "Model these Examples", br.get("examples"))
            _label(doc, "Assign practice", br.get("practice"))
        if L.get("cfu"):
            _heading(doc, "Checks for Understanding (CFU)", 11)
            _bullets(doc, L["cfu"])
        if L.get("activities"):
            _heading(doc, "Activities (do within the lesson)", 11)
            for act in L["activities"]:
                if isinstance(act, dict):
                    name = act.get("name", "")
                    tag = act.get("type", "")
                    phase = act.get("phase", "")
                    head = name + (f" ({tag})" if tag else "")
                    if phase:
                        head += f" — {phase}"
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(head).bold = True
                    if act.get("how"):
                        doc.add_paragraph(str(act["how"]))
                    if act.get("why"):
                        _label(doc, "Builds", act["why"])
                else:
                    doc.add_paragraph(str(act), style="List Bullet")
        _label(doc, "Exit Ticket", _exit_ticket_text(L.get("exit_ticket")))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def coach_summary_to_docx(summary: dict, narrative: dict,
                          framework_lens: dict | None = None) -> bytes:
    """Render the one-page Coach Summary (essentials + how-to-present narrative +
    this week's coaching-framework lens) a coach uses to lead planning."""
    doc = Document()
    _add_logo(doc)
    _heading(doc, "Coach One-Pager", 16)
    meta = doc.add_paragraph()
    m = meta.add_run(
        f"Grade {summary.get('grade_level','')} {summary.get('subject','')}  ·  "
        f"{summary.get('title','')}"
        + (f"  ·  Assessment: {summary['assessment_date']}"
           if summary.get("assessment_date") else ""))
    m.italic = True
    m.font.size = Pt(9)

    if narrative.get("big_idea"):
        _heading(doc, "Big Idea — what teachers must understand", 13)
        doc.add_paragraph(str(narrative["big_idea"]))
        if narrative.get("why_it_matters"):
            _label(doc, "Why it matters", narrative["why_it_matters"])

    if narrative.get("talking_points"):
        _heading(doc, "How to present it (your talking points)", 12)
        _bullets(doc, narrative["talking_points"], style="List Number")

    # This week's coaching-framework lens, scripted to this topic.
    if framework_lens and framework_lens.get("content"):
        c = framework_lens["content"]
        _heading(doc, f"Coaching Lens — {framework_lens.get('component_name','')} "
                 f"({framework_lens.get('week_focus','')})", 12)
        if c.get("how_it_shows_up"):
            doc.add_paragraph(str(c["how_it_shows_up"]))
        if c.get("look_fors"):
            _label(doc, "Look-fors", "")
            _bullets(doc, c["look_fors"])
        if c.get("coaching_questions"):
            _label(doc, "Coaching questions", "")
            _bullets(doc, c["coaching_questions"])
        if c.get("teacher_talking_points"):
            _label(doc, "Say this in the meeting", "")
            _bullets(doc, c["teacher_talking_points"])

    if summary.get("strategies"):
        _heading(doc, "Strategies to reinforce", 12)
        for s in summary["strategies"]:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(f"{s['name']} — ")
            r.bold = True
            p.add_run(s["what"])

    if summary.get("vocabulary"):
        _heading(doc, "Vocabulary (from the pacing guide)", 12)
        doc.add_paragraph(" · ".join(summary["vocabulary"]))

    if summary.get("sentence_frames"):
        _heading(doc, "Sentence frames / stems", 12)
        _bullets(doc, summary["sentence_frames"])

    watch = narrative.get("watch_fors") or [m["misconception"] for m in summary.get("misconceptions", [])]
    if watch:
        _heading(doc, "Watch-fors (misconceptions to flag)", 12)
        _bullets(doc, watch)
    if summary.get("misconceptions"):
        _misconception_table(
            doc, [{"misconception": m["misconception"], "example": "", "fix": m["fix"]}
                  for m in summary["misconceptions"]])

    if summary.get("level3"):
        l3 = summary["level3"]
        _heading(doc, "What a Level 3 (on-grade) looks like", 12)
        _label(doc, "Problem", l3.get("problem"))
        _label(doc, "Solution", l3.get("solution"))
        _label(doc, "Student explanation", l3.get("student_explanation"))

    if summary.get("models"):
        _heading(doc, "Worked models — how to show it (so you can model it)", 12)
        for mdl in summary["models"]:
            p = doc.add_paragraph()
            r = p.add_run(f"Lesson {mdl.get('code','')} — {mdl.get('title','')}")
            r.bold = True
            _label(doc, "Model this problem", mdl.get("problem"))
            _label(doc, "Teacher move", mdl.get("teacher_move"))
            _label(doc, "Concrete (manipulative)", mdl.get("concrete"))
            _label(doc, "Pictorial (drawing)", mdl.get("pictorial"))
            _label(doc, "Abstract (equation)", mdl.get("abstract"))

    if summary.get("lessons"):
        _heading(doc, "Lessons at a glance", 12)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(["#", "Learning goal", "Model this", "Exit ticket"]):
            table.rows[0].cells[i].text = h
        for L in summary["lessons"]:
            c = table.add_row().cells
            c[0].text = str(L.get("code", ""))
            c[1].text = str(L.get("learning_goal", "") or L.get("title", ""))
            c[2].text = str(L.get("model_focus", ""))
            c[3].text = str(L.get("exit", ""))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _cell(cell, text, size=8, bold=False, italic=False, color=None):
    """Write compact text into a table cell (first paragraph)."""
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text) if text is not None else "")
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = RGBColor(*color)
    return cell


def _grid_cell(cell, text, size=7):
    """Write a filled phase cell, bolding the leading 'Label:' on each line (e.g.
    'Collaborative strategy:', 'Strategy:', 'Connect:', 'Check:') so the box leads
    with the labeled move."""
    cell.text = ""
    p = cell.paragraphs[0]
    lines = str(text or "").split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run("\n")
        # Bold a short leading label (before the first colon) if it looks like one.
        if ":" in line and len(line.split(":", 1)[0]) <= 40:
            label, rest = line.split(":", 1)
            rb = p.add_run(label + ":")
            rb.bold = True
            rb.font.size = Pt(size)
            rr = p.add_run(rest)
            rr.font.size = Pt(size)
        else:
            r = p.add_run(line)
            r.font.size = Pt(size)
    return cell


def _vocab_line(words, cap=8):
    """'word (meaning); …' compact for the weekly one-pager strip."""
    parts = []
    for w in (words or [])[:cap]:
        term = str(w.get("word", "")).strip()
        mean = str(w.get("meaning", "")).strip()
        parts.append(f"{term} ({mean})" if mean else term)
    return " · ".join(parts)


def template_to_docx(t: dict, filled: bool = False) -> bytes:
    """Render the WEEKLY teacher walkout: a one-page landscape grid — 5 lesson
    slots (each with a DATE blank, not a weekday) across, the gradual-release
    phases down — with a Tier 2 / Tier 3 vocabulary strip. Phase cells are blank
    for the teacher to fill; when `filled` is True they carry a worked example
    from the guide so the coach can show what a completed plan looks like."""
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, m, Inches(0.4))

    # Compact header.
    h = doc.add_paragraph()
    hr = h.add_run("Weekly Collaborative Planning — My Lesson Plan"
                   + ("  (EXAMPLE)" if filled else ""))
    hr.bold = True
    hr.font.size = Pt(14)
    hr.font.color.rgb = RGBColor(0x38, 0x60, 0x1F)
    meta = doc.add_paragraph()
    mr = meta.add_run(
        f"Grade {t.get('grade_level','')} {t.get('subject','')}  ·  {t.get('title','')}"
        f"  ·  Standard(s): "
        + ", ".join(b.get("code", "") for b in t.get("benchmarks", []))
        + "      Teacher: __________________   Week of: __________________")
    mr.font.size = Pt(9)
    if t.get("learning_goal"):
        lg = doc.add_paragraph()
        lgr = lg.add_run(f"Learning goal: {t['learning_goal']}")
        lgr.font.size = Pt(9)
        lgr.italic = True

    # Vocabulary strip (Tier 2 / Tier 3) — the week's focus.
    voc = t.get("vocabulary", {}) or {}
    vt = doc.add_table(rows=2, cols=2)
    vt.style = "Light Grid Accent 1"
    _cell(vt.rows[0].cells[0], "Power Words — academic (words in the questions)", 8,
          bold=True, color=(0x2E, 0x86, 0xC1))
    _cell(vt.rows[0].cells[1], _vocab_line(voc.get("tier2")), 8)
    _cell(vt.rows[1].cells[0], "Tier 3 — subject-specific math words", 8,
          bold=True, color=(0xE6, 0x7E, 0x22))
    _cell(vt.rows[1].cells[1], _vocab_line(voc.get("tier3")), 8)
    vt.columns[0].width = Inches(2.4)
    vt.columns[1].width = Inches(7.6)

    doc.add_paragraph().add_run("").font.size = Pt(2)

    # The weekly grid: rows = section labels, cols = the 5 days.
    days = t.get("days", [])
    phases = t.get("phases", [])
    ncols = 1 + len(days)
    # Row plan: header, learning goal, one row per phase, exit check.
    row_defs = [("Learning goal / focus", "goal")]
    for ph in phases:
        row_defs.append((f"{ph['gradual_release']} ({ph['aces']})", ph["key"]))
    row_defs.append(("Activities (in-lesson)", "activities"))
    row_defs.append(("Exit check (CFU)", "exit"))

    table = doc.add_table(rows=len(row_defs) + 1, cols=ncols)
    table.style = "Table Grid"

    # Header row: blank corner + each lesson slot with a DATE blank (no weekday).
    _cell(table.rows[0].cells[0], "Date →", 8, bold=True)
    for j, d in enumerate(days):
        lesson = f"{d.get('lesson_code','')} {d.get('title','')}".strip()
        head = (f"Lesson {lesson}" if lesson else f"Lesson {d.get('slot','')}")
        head += "\nDate: ____________"
        _cell(table.rows[0].cells[j + 1], head, 8, bold=True,
              color=(0x38, 0x60, 0x1F))

    # Body rows.
    phase_by_key = {ph["key"]: ph for ph in phases}
    for i, (label, key) in enumerate(row_defs):
        r = table.rows[i + 1]
        # Row label cell (with the "what to plan" reminder for phases).
        lc = r.cells[0]
        lc.text = ""
        p = lc.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(8)
        if key in phase_by_key:
            sub = p.add_run("\nwhat you do · questions · what students do")
            sub.italic = True
            sub.font.size = Pt(6.5)
            sub.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        # Day cells.
        for j, d in enumerate(days):
            cell = r.cells[j + 1]
            if key == "goal":
                _cell(cell, d.get("learning_goal", "") or "", 7.5)
            elif key == "exit":
                _cell(cell, d.get("exit", "") or "", 7.5)
            elif key == "activities":
                # Always show the lesson's activities (they're a menu to pick from).
                _cell(cell, d.get("activities", "") or "", 7)
            elif filled:
                _grid_cell(cell, (d.get("phase_example", {}) or {}).get(key, ""), 7)
            else:
                _cell(cell, "", 8)  # blank for the teacher to plan

    # Column widths: label column narrow, day columns share the rest.
    table.columns[0].width = Inches(1.5)
    for j in range(1, ncols):
        table.columns[j].width = Inches(8.5 / max(1, len(days)))

    # Compact footer: sentence frames + a misconception to plan for.
    if t.get("sentence_frames"):
        fp = doc.add_paragraph()
        fr = fp.add_run("Sentence frames: "
                        + "  |  ".join(f"“{s}”" for s in t["sentence_frames"][:3]))
        fr.font.size = Pt(8)
    if t.get("misconception"):
        mc = t["misconception"]
        mp = doc.add_paragraph()
        mr2 = mp.add_run(
            f"Misconception to watch: {mc.get('misconception','')}  →  "
            f"Fix: {mc.get('fix','')}")
        mr2.font.size = Pt(8)
        mr2.font.color.rgb = RGBColor(0xB0, 0x30, 0x30)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _di_phase(doc, label, ph):
    """One gradual-release phase inside a teacher-led (TLC) reteach session."""
    if not isinstance(ph, dict):
        return
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    if ph.get("problem"):
        p.add_run(str(ph["problem"]))
    for line in ph.get("say", []) or []:
        doc.add_paragraph(f"“{line}”", style="List Bullet")
    if ph.get("do"):
        _label(doc, "Teacher", ph["do"])
    if ph.get("answer"):
        _label(doc, "Answer", ph["answer"])


def di_packets_to_docx(packet: dict):
    """Render the three-tier DI packet (Intensive / Cusp / Strategic) as a
    printable Word doc: each tier's rotation, teacher-led reteach, independent
    practice, and OPM progress check. Returns a python-docx Document."""
    doc = Document()
    _add_logo(doc)
    _heading(doc, f"Differentiated Instruction Packets — {packet.get('standard','')}", 16)
    _label(doc, "Grade", packet.get("grade_level", ""))
    _label(doc, "Benchmark", packet.get("description", ""))
    doc.add_paragraph(
        "DI Rotation stations: i-Ready · TLC (teacher-led) · IXL/Skill Trainer/"
        "Independent Practice · OPM · Data Chat.").italic = True

    for t in packet.get("tiers", []) or []:
        doc.add_page_break()
        _heading(doc, f"{t.get('tier','')} — {t.get('band','')} "
                      f"({'★' * int(t.get('stars', 0))})", 14,
                 color=(0x2E, 0x86, 0xC1))
        _label(doc, "Students in this group", t.get("student_count", 0))
        _label(doc, "Teacher-led (TLC) sessions in the rotation", t.get("tlc_sessions", 0))
        if t.get("rotation"):
            days = "  ".join(f"D{i+1}:{st}" for i, st in enumerate(t["rotation"]))
            _label(doc, "7-day rotation", days)
        if t.get("focus"):
            _label(doc, "Focus", t["focus"])

        for sess in t.get("teacher_led", []) or []:
            _heading(doc, f"Teacher-Led (TLC) Session {sess.get('session','')}: "
                          f"{sess.get('title','')}", 12, color=(0x38, 0x60, 0x1F))
            _di_phase(doc, "I Do (model)", sess.get("i_do"))
            _di_phase(doc, "We Do (guided)", sess.get("we_do"))
            _di_phase(doc, "You Do (independent)", sess.get("you_do"))

        ip = t.get("independent_practice") or []
        if ip:
            _heading(doc, "Independent Practice (IXL / Skill Trainer station)", 12)
            for q in ip:
                p = doc.add_paragraph(style="List Number")
                p.add_run(str(q.get("problem", "")))
                if q.get("answer"):
                    a = p.add_run(f"   (answer: {q['answer']})")
                    a.italic = True

        opm = t.get("opm") or []
        if opm:
            _heading(doc, "OPM — Progress Monitoring (did they grow?)", 12,
                     color=(0xB0, 0x30, 0x30))
            for q in opm:
                p = doc.add_paragraph(style="List Number")
                p.add_run(str(q.get("problem", "")))
                if q.get("answer"):
                    a = p.add_run(f"   (answer: {q['answer']})")
                    a.italic = True

        if t.get("students"):
            names = ", ".join(s.get("student_name", "") for s in t["students"])
            _label(doc, "Group roster", names)

    return doc
