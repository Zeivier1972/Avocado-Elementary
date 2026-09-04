"""Extract plain text from an uploaded document (pacing guide, resource, etc.)
so the planning-guide generator can be grounded in what the coach uploaded.

Supports PDF (pypdf), Word (.docx), Excel (.xlsx), and plain text/CSV. Returns
("", reason) when the file has no extractable text (e.g. a scanned/image-only
PDF) so the caller can explain it to the user."""
from __future__ import annotations

import io


def extract_document_text(filename: str, content_type: str, data: bytes) -> tuple[str, str]:
    name = (filename or "").lower()
    ct = (content_type or "").lower()

    # PDF
    if name.endswith(".pdf") or "pdf" in ct:
        return _from_pdf(data)
    # Word
    if name.endswith(".docx") or "wordprocessingml" in ct:
        return _from_docx(data)
    # Excel
    if name.endswith((".xlsx", ".xlsm")) or "spreadsheetml" in ct:
        return _from_xlsx(data)
    # Plain text / CSV
    if name.endswith((".txt", ".csv", ".md")) or ct.startswith("text/"):
        try:
            return data.decode("utf-8-sig", errors="replace"), ""
        except Exception as e:  # pragma: no cover
            return "", f"could not decode text file: {e}"
    # Fallback: try PDF, then text
    txt, _ = _from_pdf(data)
    if txt:
        return txt, ""
    try:
        return data.decode("utf-8-sig", errors="replace"), ""
    except Exception:
        return "", "unsupported file type — upload a PDF, Word, Excel, or text pacing guide"


def _from_pdf(data: bytes) -> tuple[str, str]:
    """Read a PDF's text. Three passes, best-effort, in order of cost:
    1) pypdf (fast, works on normal text PDFs);
    2) PyMuPDF (recovers a text layer pypdf misses in many 'looks-scanned' PDFs);
    3) Tesseract OCR page-by-page (for truly scanned/image-only PDFs like a
       photographed textbook). OCR is slower but runs in the background job."""
    text = ""
    # Pass 1 — pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        parts = [(p.extract_text() or "") for p in reader.pages]
        text = "\n".join(parts).strip()
    except Exception:
        text = ""
    if len(text) >= 200:
        return text, ""

    # Pass 2 — PyMuPDF text layer
    mupdf_text = _pymupdf_text(data)
    if len(mupdf_text) >= 200:
        return mupdf_text, ""
    text = text or mupdf_text

    # Pass 3 — OCR (scanned pages)
    ocr_text, ocr_reason = _ocr_pdf(data)
    if len(ocr_text) >= 40:
        return ocr_text, ""

    if text:
        return text, ""
    return "", (ocr_reason or "this PDF has no selectable text (it looks "
                "scanned/image-only) and OCR could not read it")


def _pymupdf_text(data: bytes) -> str:
    try:
        try:
            import pymupdf as fitz  # newer name
        except ImportError:
            import fitz  # older name
    except Exception:
        return ""
    try:
        doc = fitz.open(stream=data, filetype="pdf")
        parts = [page.get_text("text") or "" for page in doc]
        doc.close()
        return "\n".join(parts).strip()
    except Exception:
        return ""


_OCR_MAX_PAGES = 80  # a full textbook chapter; keeps OCR bounded


def _ocr_pdf(data: bytes) -> tuple[str, str]:
    """OCR a scanned PDF with Tesseract, rendering each page via PyMuPDF (no
    poppler needed). Returns (text, reason)."""
    try:
        try:
            import pymupdf as fitz
        except ImportError:
            import fitz
    except Exception:
        return "", "OCR needs PyMuPDF (not installed)"
    try:
        import pytesseract
        from PIL import Image
    except Exception:
        return "", "OCR needs pytesseract + Pillow (not installed)"
    try:
        doc = fitz.open(stream=data, filetype="pdf")
    except Exception as e:
        return "", f"could not open PDF for OCR: {e}"
    parts: list[str] = []
    # ~200 DPI is enough for printed textbook text and keeps each page fast.
    zoom = fitz.Matrix(200 / 72, 200 / 72)
    try:
        for i, page in enumerate(doc):
            if i >= _OCR_MAX_PAGES:
                break
            try:
                pix = page.get_pixmap(matrix=zoom)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                parts.append(pytesseract.image_to_string(img) or "")
            except Exception:
                continue
    finally:
        doc.close()
    text = "\n".join(parts).strip()
    if not text:
        return "", ("OCR ran but found no readable text — the scan may be too "
                    "low-quality")
    return text, ""


def _from_docx(data: bytes) -> tuple[str, str]:
    try:
        from docx import Document
    except ImportError:
        return "", "Word support not installed (python-docx)"
    try:
        doc = Document(io.BytesIO(data))
        lines = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    lines.append(" | ".join(cells))
        text = "\n".join(lines).strip()
        return (text, "") if text else ("", "the Word document has no text")
    except Exception as e:
        return "", f"could not read Word document: {e}"


def _from_xlsx(data: bytes) -> tuple[str, str]:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "", "Excel support not installed (openpyxl)"
    try:
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        lines = []
        for ws in wb.worksheets:
            lines.append(f"# Sheet: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c not in (None, "")]
                if cells:
                    lines.append(" | ".join(cells))
        text = "\n".join(lines).strip()
        return (text, "") if text else ("", "the spreadsheet has no data")
    except Exception as e:
        return "", f"could not read spreadsheet: {e}"
